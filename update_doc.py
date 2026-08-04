from docx import Document

doc = Document('CodeFundamentals_Work_Summary.docx')

doc.add_paragraph('')

p = doc.add_paragraph()
run = p.add_run('Session coverage: 5:00 PM onwards | Date: August 4, 2026')
run.bold = True
p.add_run('\nPlatform: Flask + Bootstrap 5.3.3 web application | Repository: github.com/christianrenz210/learnfUNDAMENTALS')

doc.add_heading('11. Advanced Fundamentals Lessons (4 Lessons)', level=1)
items = [
    'Added 4 new lessons to the Advanced Fundamentals section: Introduction to GitHub (id=31), Web Deployment Fundamentals (id=30), Deploying with Render (id=32), Deploying with Netlify (id=33).',
    'GitHub lesson: Shell code examples covering git init, commit, push, pull, branch, merge, and pull requests. Video: tRZGeaHPoaw.',
    'Web Deployment Fundamentals: Laravel + React code examples covering deployment concepts, environment variables, CI/CD, and hosting platforms. Videos: m-QO9Qp_wRQ, SuZBpX7Y7EA.',
    'Deploying with Render: Comprehensive 40-minute lesson with 20 key points, 12 quiz questions, and detailed step-by-step code examples for Laravel, React, databases, custom domains, and monitoring. Video: srudWEWKPv0.',
    'Deploying with Netlify: 35-minute lesson with 20 key points, 12 quiz questions, covering netlify.toml, forms, functions, CLI, drag-and-drop, and custom domains. Videos: IH95RG9Y6L4, YUtNzwxOVYY.',
    'Updated EREN_SYSTEM_PROMPT, get_video_reply(), assistant_reply(), and GENERAL_KNOWLEDGE in app.py for all new lessons.',
    'Fixed invalid Render video ID and certificate route accidentally deleted during API endpoint addition.',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('12. Admin Online/Offline User Tracking', level=1)
items = [
    'Added last_seen column to User model with auto-migration on startup (supports both SQLite and PostgreSQL).',
    'Added before_request hook to update last_seen on every authenticated user request.',
    'Admin Dashboard: New Online Now stat card with auto-refresh every 30 seconds via JavaScript.',
    'Admin Users page: Added Status column (Online/Offline badge), Last Active column, and online/offline count in header. Auto-refreshes every 30 seconds.',
    'API endpoints: /admin/api/online-users (returns all users with online status) and /admin/api/stats (returns dashboard stats).',
    'User is considered online if last_seen is within 5 minutes of current time.',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('13. Admin User Progress Management', level=1)
items = [
    'New admin route: /admin/user/<id>/progress - displays full user progress with lesson completion, quiz scores, and lab completions.',
    'Reset Lesson: Admin can reset individual lesson completion, quiz score, and lab completion with confirmation prompt.',
    'Reset Section Quiz: Admin can reset long quiz scores for specific sections.',
    'Reset All Progress: Admin can delete ALL user progress (lessons, quizzes, section quizzes, labs) with double confirmation.',
    'Admin Users page: Added Progress action button for each user.',
    'All admin actions are logged to ActivityLog for audit trail.',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('14. Landing Page Update', level=1)
items = [
    'Hero section updated to mention Web Deployment lessons.',
    'Added 3 new feature cards: Learn to Deploy (GitHub, Render, Netlify), AI Assistant (E.R.E.N), Leaderboards and Certificates.',
    'Added Advanced Fundamentals section to What You Will Learn with 4 lessons.',
    'Updated lesson counts: Fundamentals (11), Advanced Fundamentals (4), Database (8), SQL Injection (2).',
    'Updated How It Works section: 25+ lessons in four sections.',
    'Removed Coming Soon badges from SQL Injection section (only showing active lessons).',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('15. Feedback System', level=1)
items = [
    'New Feedback model in database: id, user_id, name, email, category, message, admin_reply, is_read, created_at, replied_at.',
    'Public feedback page (/feedback): Viewable by anyone without login. Shows all feedbacks with user avatars.',
    'Submit form: Only logged-in users can submit feedback. Non-authenticated users see a login/register prompt.',
    'User avatars: Profile picture displayed next to feedback, falls back to initials if no avatar.',
    'Admin replies: Visible publicly with green highlight and admin badge.',
    'Admin panel (/admin/feedback): View all feedbacks, mark as read, reply to feedback, delete feedback.',
    'Feedback categories: General, Bug Report, Feature Request, Lesson Content, Quiz Issue, Other.',
    'Added Feedback link to navbar for both authenticated and unauthenticated users.',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('16. FAQ Page Fix', level=1)
items = [
    'Updated lesson count from 21 to 25.',
    'Updated section description: now mentions 4 sections including Advanced Fundamentals.',
    'Updated languages: now includes Python, C++, Java, SQL, and Web Deployment.',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('17. Commits Pushed', level=1)
items = [
    '7363deb - Advanced fundamentals + admin tracking + user progress management',
    '8b0dfa6 - last_seen column migration fix for Render PostgreSQL',
    '6456354 - Landing page update with new features',
    '467e8c6 - Feedback page + FAQ fix + admin feedback panel',
    'd705dc9 - Public feedback page + user avatars',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.save('CodeFundamentals_Work_Summary.docx')
print('Document updated successfully!')
print('Total paragraphs:', len(doc.paragraphs))
